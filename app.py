import streamlit as st
from google import genai
from supabase import create_client
import json
from fpdf import FPDF
import tempfile
import secrets as pysecrets
import arabic_reshaper
from bidi.algorithm import get_display
from streamlit_markmap import markmap
from streamlit_option_menu import option_menu

URDU_FONT_PATH = "fonts/NotoSansArabic.ttf"

def is_urdu_text(text):
    return any('\u0600' <= ch <= '\u06FF' or '\u0750' <= ch <= '\u077F'
               or '\uFB50' <= ch <= '\uFDFF' or '\uFE70' <= ch <= '\uFEFF' for ch in str(text))

def shape_urdu(text):
    reshaped = arabic_reshaper.reshape(str(text))
    return get_display(reshaped)

st.set_page_config(page_title="Summer Task Generator", page_icon="📚", layout="centered")

st.markdown("""
<style>
    .question-card {
        background-color: #F1F8F4;
        border: 1px solid #D7EAD9;
        border-radius: 10px;
        padding: 16px 20px;
        margin-bottom: 14px;
    }
    .question-card .q-number {
        color: #2E7D32;
        font-weight: 700;
        font-size: 0.85rem;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    .question-card .q-text {
        font-size: 1.05rem;
        font-weight: 600;
        margin: 4px 0 8px 0;
    }
    .question-card .q-option {
        padding: 2px 0 2px 12px;
        color: #333;
    }
    .question-card .q-answer {
        margin-top: 8px;
        color: #2E7D32;
        font-style: italic;
        font-size: 0.92rem;
    }
    .app-header {
        text-align: center;
        padding: 10px 0 20px 0;
    }
    .app-header h1 {
        margin-bottom: 0;
    }
    .app-header p {
        color: #667;
        margin-top: 4px;
    }
    .stat-card {
        background-color: #F1F8F4;
        border: 1px solid #D7EAD9;
        border-radius: 12px;
        padding: 18px 20px;
        text-align: center;
    }
    .stat-card .stat-number {
        font-size: 2rem;
        font-weight: 700;
        color: #2E7D32;
        line-height: 1.1;
    }
    .stat-card .stat-label {
        color: #556;
        font-size: 0.85rem;
        margin-top: 4px;
    }
    .dash-welcome {
        padding: 4px 0 18px 0;
    }
    .dash-welcome h2 {
        margin-bottom: 2px;
    }
    .dash-welcome p {
        color: #667;
    }
    .recent-task-row {
        border-bottom: 1px solid #EEE;
        padding: 10px 2px;
    }
    section[data-testid="stSidebar"] {
        border-right: 1px solid #EEE;
    }
</style>
""", unsafe_allow_html=True)


def render_question_card(i, q, show_answer=True):
    """Renders one question as a styled card. Works for st.markdown with HTML."""
    html = f'<div class="question-card">'
    html += f'<div class="q-number">Question {i}</div>'
    html += f'<div class="q-text">{q["question"]}</div>'
    if q["type"] == "mcq":
        for idx, opt in enumerate(q["options"]):
            html += f'<div class="q-option">{chr(65 + idx)}. {opt}</div>'
    if show_answer:
        html += f'<div class="q-answer">Answer: {q["answer"]}</div>'
    html += '</div>'
    st.markdown(html, unsafe_allow_html=True)


def render_table_card(q, show_answer=True):
    """Renders a Gender/Number grammar exercise as a right-to-left bordered table."""
    st.markdown(f"**{q['title']}**")
    rows = ""
    for p in q["pairs"]:
        ans_cell = p["answer"] if show_answer else "______"
        rows += (f'<tr><td style="padding:8px 14px;border:1px solid #D7EAD9;text-align:right;">{p["word"]}</td>'
                 f'<td style="padding:8px 14px;border:1px solid #D7EAD9;text-align:right;">{ans_cell}</td></tr>')
    html = (f'<table dir="rtl" style="width:100%;border-collapse:collapse;font-size:1.1rem;margin-bottom:14px;">'
            f'<tr><th style="padding:8px 14px;border:1px solid #D7EAD9;background:#F1F8F4;">لفظ</th>'
            f'<th style="padding:8px 14px;border:1px solid #D7EAD9;background:#F1F8F4;">جواب</th></tr>{rows}</table>')
    st.markdown(html, unsafe_allow_html=True)

# --- Connections ---
api_key = st.secrets["GEMINI_API_KEY"]
client = genai.Client(api_key=api_key)

supabase_url = st.secrets["SUPABASE_URL"]
supabase_key = st.secrets["SUPABASE_KEY"]
supabase = create_client(supabase_url, supabase_key)  # used for auth (signup/login/logout)

service_key = st.secrets["SUPABASE_SERVICE_KEY"]
db = create_client(supabase_url, service_key)  # used for table reads/writes (server-side only, bypasses RLS)

app_url = st.secrets.get("APP_URL", "")

CLASS_OPTIONS = ["Nursery", "KG", "Class 1", "Class 2", "Class 3", "Class 4", "Class 5",
                  "Class 6", "Class 7", "Class 8", "Class 9", "Class 10", "Class 11", "Class 12"]

ALL_QUESTION_TYPES = ["mcq", "true_false", "fill_blank", "theory", "theory_short", "theory_long", "math", "mind_map",
                       "word_meaning", "synonym_antonym", "sentence_making", "tenses",
                       "urdu_grammar", "translation", "gender_table", "number_table"]

URDU_SCRIPT_TYPES = {"urdu_grammar", "translation", "gender_table", "number_table"}

TYPE_LABELS = {
    "mcq": "MCQ",
    "true_false": "True/False",
    "fill_blank": "Fill in the Blank",
    "theory": "Theory (medium length)",
    "theory_short": "Theory - Short Answer",
    "theory_long": "Theory - Long Answer",
    "math": "Math",
    "mind_map": "Mind Map",
    "word_meaning": "Word Meaning",
    "synonym_antonym": "Synonym/Antonym",
    "sentence_making": "Sentence Making",
    "tenses": "Tenses",
    "urdu_grammar": "Urdu Grammar (Qawaid)",
    "translation": "Translation",
    "gender_table": "Gender Table (مذکر/مؤنث)",
    "number_table": "Singular/Plural Table (واحد/جمع)"
}

# --- Auth helpers ---
def sign_up(email, password):
    return supabase.auth.sign_up({"email": email, "password": password})

def sign_in(email, password):
    return supabase.auth.sign_in_with_password({"email": email, "password": password})

def sign_out():
    supabase.auth.sign_out()
    st.session_state.clear()

def reset_password(email):
    supabase.auth.reset_password_for_email(email)


# --- Usage quota (protects the shared free-tier Gemini API key across all schools) ---
def check_quota(school_id, daily_limit):
    """Returns (allowed, used_today, limit). Counts one unit per 'Generate' click, not per question."""
    today_start = __import__("datetime").datetime.utcnow().strftime("%Y-%m-%dT00:00:00")
    used_today = len(
        db.table("usage_log").select("id").eq("school_id", school_id).gte("created_at", today_start).execute().data
    )
    return used_today < daily_limit, used_today, daily_limit


def log_usage(school_id, teacher_id):
    try:
        db.table("usage_log").insert({"school_id": school_id, "teacher_id": teacher_id}).execute()
    except Exception:
        pass  # usage logging should never block the actual generation


# --- Question generation ---
def _is_valid_question(q):
    """Structural validation for one generated question, before it ever reaches a student."""
    if not isinstance(q, dict) or not q.get("type"):
        return False
    qtype = q["type"]
    if qtype == "mind_map":
        return bool(q.get("title")) and isinstance(q.get("children"), list) and len(q["children"]) > 0
    if qtype in ("gender_table", "number_table"):
        pairs = q.get("pairs")
        return bool(q.get("title")) and isinstance(pairs, list) and len(pairs) > 0 and all(
            isinstance(p, dict) and str(p.get("word", "")).strip() and str(p.get("answer", "")).strip() for p in pairs
        )
    if not str(q.get("question", "")).strip() or not str(q.get("answer", "")).strip():
        return False
    if qtype == "mcq":
        options = q.get("options")
        if not isinstance(options, list) or len(options) < 2:
            return False
        cleaned = [str(o).strip().lower() for o in options]
        if len(set(cleaned)) != len(cleaned):
            return False  # duplicate options
        if str(q["answer"]).strip().lower() not in cleaned:
            return False  # answer must actually be one of the listed options
    return True


def _dedupe_questions(q_list):
    seen = set()
    result = []
    for q in q_list:
        key = str(q.get("question") or q.get("title") or "").strip().lower()
        if key and key in seen:
            continue
        if key:
            seen.add(key)
        result.append(q)
    return result


def generate_questions(class_level, subject, topic, question_type, count, difficulty="medium",
                        source_text=None, chapter_focus=None, answer_length="medium"):
    # theory_short / theory_long are just "theory" with a forced answer length, so a teacher
    # can pick both in the same task (each with its own count) via the question type list.
    effective_type = question_type
    if question_type == "theory_short":
        effective_type = "theory"
        answer_length = "short"
    elif question_type == "theory_long":
        effective_type = "theory"
        answer_length = "long"

    format_instructions = {
        "mcq": '''[{"type": "mcq", "question": "...", "options": ["...", "...", "...", "..."], "answer": "..."}]''',
        "true_false": '''[{"type": "true_false", "question": "...", "answer": "True or False"}]''',
        "fill_blank": '''[{"type": "fill_blank", "question": "sentence with ____ blank", "answer": "..."}]''',
        "theory": '''[{"type": "theory", "question": "...", "answer": "model answer text"}]''',
        "math": '''[{"type": "math", "question": "...", "answer": "final answer", "solution": "step by step solution"}]''',
        "mind_map": '''[{"type": "mind_map", "title": "Central Topic", "children": [{"title": "Main Branch 1", "children": [{"title": "Sub-point", "children": []}]}, {"title": "Main Branch 2", "children": []}]}]''',
        "word_meaning": '''[{"type": "word_meaning", "question": "What is the meaning of the word '...'?", "answer": "..."}]''',
        "synonym_antonym": '''[{"type": "synonym_antonym", "question": "Give a synonym and an antonym of the word '...'.", "answer": "Synonym: ...   Antonym: ..."}]''',
        "sentence_making": '''[{"type": "sentence_making", "question": "Use the word '...' in a meaningful sentence.", "answer": "example sentence using the word"}]''',
        "tenses": '''[{"type": "tenses", "question": "Change the following sentence's tense as instructed: '...'", "answer": "corrected sentence"}]''',
        "urdu_grammar": '''[{"type": "urdu_grammar", "question": "...(Urdu grammar/qawaid question in Urdu script)...", "answer": "..."}]''',
        "translation": '''[{"type": "translation", "question": "Translate the following into Urdu/English as instructed: '...'", "answer": "..."}]''',
        "gender_table": '''[{"type": "gender_table", "title": "\u0645\u0630\u06a9\u0631 \u0633\u06d2 \u0645\u0624\u0646\u062b \u0628\u0646\u0627\u0626\u06cc\u06ba", "pairs": [{"word": "\u0644\u0691\u06a9\u0627", "answer": "\u0644\u0691\u06a9\u06cc"}]}]''',
        "number_table": '''[{"type": "number_table", "title": "\u0648\u0627\u062d\u062f \u0633\u06d2 \u062c\u0645\u0639 \u0628\u0646\u0627\u0626\u06cc\u06ba", "pairs": [{"word": "\u06a9\u062a\u0627\u0628", "answer": "\u06a9\u062a\u0627\u0628\u06cc\u06ba"}]}]'''
    }

    if source_text:
        # Trim very long notes to keep prompt size reasonable
        trimmed = source_text[:15000]
        focus_line = f'\nFocus specifically on the chapter/topic "{chapter_focus}" within this material \u2014 ignore unrelated sections if present.\n' if chapter_focus else ""
        source_instruction = f"""
Base the questions STRICTLY on the following notes/chapter content provided by the teacher.
Do not use outside knowledge beyond what's in this text.
{focus_line}
--- NOTES START ---
{trimmed}
--- NOTES END ---
"""
    else:
        source_instruction = f'Base the questions on general knowledge of the topic "{topic}".'

    if effective_type == "mind_map":
        count_instruction = f"with about {count} main branches, each with 2-4 sub-points"
    elif effective_type in ("gender_table", "number_table"):
        count_instruction = f"with {count} word pairs"
    else:
        count_instruction = f"{count} separate questions"

    length_instruction = {
        "short": "Keep answers concise \u2014 1-2 sentences for theory-type answers, brief phrases otherwise. For theory/essay-style question types, phrase the question as a short-answer question.",
        "long": "Provide detailed, well-explained answers \u2014 aim for a full, well-structured paragraph with reasoning for theory-type answers. For theory/essay-style question types, phrase the question as a long-answer / detailed-explanation question.",
        "medium": "Keep answers reasonably thorough but not excessively long."
    }.get(answer_length, "")

    quality_instruction = """
IMPORTANT \u2014 question quality rules:
- Every question must test genuine conceptual understanding: a specific mechanism, cause-and-effect relationship, comparison, or application \u2014 not a vague or generic fact a student could guess without understanding the material.
- Avoid overly broad or ambiguous phrasing (e.g. do not simply ask "What is X?" if a more specific, reasoning-based question is possible).
- Each question should have exactly one clear, unambiguous correct answer.
"""

    prompt = f"""
You are a question generator for a school summer task app.

Generate a {effective_type} for {class_level} {subject}, at {difficulty} difficulty, {count_instruction}.

{quality_instruction}
{length_instruction}

{source_instruction}

Return ONLY valid JSON, no extra text, no markdown code fences, in this exact format:

{format_instructions[effective_type]}
"""
    import time
    last_error = None
    parsed = None
    for attempt in range(3):
        try:
            response = client.models.generate_content(model="gemini-3.5-flash", contents=prompt)
            text = response.text.strip()
            if text.startswith("```"):
                text = text.strip("`")
                text = text.replace("json", "", 1).strip()
            parsed = json.loads(text)
            break
        except json.JSONDecodeError as e:
            last_error = e
            continue  # malformed JSON from the model \u2014 worth a fresh attempt
        except Exception as e:
            last_error = e
            error_name = type(e).__name__
            if "ServerError" in error_name or "503" in str(e) or "500" in str(e) or "overloaded" in str(e).lower():
                time.sleep(2 * (attempt + 1))  # brief backoff before retrying
                continue
            else:
                raise  # not a transient/parsing issue, don't retry

    if parsed is None:
        raise RuntimeError(
            "The AI service (Gemini) is temporarily overloaded or returned an unusable response after several attempts. "
            "This is usually resolved within a minute or two \u2014 please wait briefly and try again."
        ) from last_error

    # Drop structurally broken questions (missing fields, MCQ answer not among its own options, etc.)
    # and remove exact duplicate questions within this batch.
    valid = _dedupe_questions([q for q in parsed if _is_valid_question(q)])

    # For standard list-style types, if quality filtering left us short of what was asked for,
    # make one best-effort top-up request for just the shortfall rather than silently under-delivering.
    if effective_type not in ("mind_map", "gender_table", "number_table") and len(valid) < count:
        shortfall = count - len(valid)
        topup_prompt = prompt.replace(f"{count} separate questions", f"{shortfall} separate questions")
        try:
            response2 = client.models.generate_content(model="gemini-3.5-flash", contents=topup_prompt)
            text2 = response2.text.strip()
            if text2.startswith("```"):
                text2 = text2.strip("`")
                text2 = text2.replace("json", "", 1).strip()
            extra = json.loads(text2)
            extra_valid = [q for q in extra if _is_valid_question(q)]
            valid = _dedupe_questions(valid + extra_valid)
        except Exception:
            pass  # top-up is best-effort \u2014 keep whatever already passed validation

    return valid[:count] if effective_type not in ("mind_map", "gender_table", "number_table") else valid


def get_pdf_page_count(uploaded_file):
    if uploaded_file.name.lower().endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(uploaded_file)
        return len(reader.pages)
    return None


def extract_text_from_file(uploaded_file, page_start=None, page_end=None):
    """Extracts text from an uploaded PDF, DOCX, or TXT file. page_start/page_end (1-indexed, inclusive) only apply to PDFs."""
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(uploaded_file)
        total_pages = len(reader.pages)
        start = (page_start - 1) if page_start else 0
        end = page_end if page_end else total_pages
        start = max(0, min(start, total_pages - 1))
        end = max(start + 1, min(end, total_pages))
        text = ""
        for page in reader.pages[start:end]:
            text += page.extract_text() or ""
        return text

    elif name.endswith(".docx"):
        from docx import Document
        doc = Document(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs)

    elif name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8")

    else:
        return ""


def mindmap_to_markdown(node, level=1):
    """Converts a mind map node (title + children) into markdown outline for markmap."""
    md = ("#" * level) + " " + node["title"] + "\n" if level == 1 else ("  " * (level - 1)) + "- " + node["title"] + "\n"
    for child in node.get("children", []):
        md += mindmap_to_markdown(child, level + 1)
    return md

def mindmap_to_pdf_lines(node, level=0):
    """Converts a mind map node into indented text lines for the PDF."""
    lines = [("    " * level) + ("- " if level > 0 else "") + node["title"]]
    for child in node.get("children", []):
        lines += mindmap_to_pdf_lines(child, level + 1)
    return lines


# --- PDF generation ---
def clean_text(text):
    return text.encode("latin-1", "replace").decode("latin-1")

def hex_to_rgb(hex_color):
    hex_color = (hex_color or "2E7D32").lstrip("#")
    if len(hex_color) != 6:
        hex_color = "2E7D32"
    return tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))


def lighten(rgb, factor=0.88):
    return tuple(int(c + (255 - c) * factor) for c in rgb)


class StyledPDF(FPDF):
    def __init__(self, school_name, logo_path, class_level, subject, topic, include_answers, brand_color="2E7D32",
                 side_margin=15, margin_mode="uniform", binding_extra=10, show_margin_lines=False, bottom_margin=22):
        super().__init__()
        self.school_name = school_name
        self.logo_path = logo_path
        self.class_level = class_level
        self.subject = subject
        self.topic = topic
        self.include_answers = include_answers
        self.brand_rgb = hex_to_rgb(brand_color)
        self.side_margin = side_margin
        self.margin_mode = margin_mode
        self.binding_extra = binding_extra
        self.show_margin_lines = show_margin_lines
        self.bottom_margin = bottom_margin

    def header(self):
        # Set this page's left/right margins based on the chosen layout, before drawing anything.
        # "uniform" = same margin every side/page. "single" = extra binding margin always on the left
        # (for one-sided printing, e.g. hole-punching or stapling). "double" = binding margin alternates
        # sides on odd/even pages, so facing pages line up correctly when printed double-sided.
        if self.margin_mode == "single":
            self.set_left_margin(self.side_margin + self.binding_extra)
            self.set_right_margin(self.side_margin)
        elif self.margin_mode == "double":
            if self.page_no() % 2 == 1:
                self.set_left_margin(self.side_margin + self.binding_extra)
                self.set_right_margin(self.side_margin)
            else:
                self.set_left_margin(self.side_margin)
                self.set_right_margin(self.side_margin + self.binding_extra)
        else:
            self.set_left_margin(self.side_margin)
            self.set_right_margin(self.side_margin)

        if self.show_margin_lines:
            self.set_draw_color(190, 190, 190)
            self.set_line_width(0.2)
            guide_top = 4
            guide_bottom = 297 - self.bottom_margin
            self.rect(self.l_margin, guide_top, 210 - self.l_margin - self.r_margin, guide_bottom - guide_top)

        # Brand-colored banner across the top of every page
        self.set_fill_color(*self.brand_rgb)
        self.rect(0, 0, 210, 28, style="F")

        if self.logo_path:
            self.image(self.logo_path, x=8, y=5, w=18)
            text_x = 30
        else:
            text_x = 10

        self.set_xy(text_x, 6)
        self.set_text_color(255, 255, 255)
        self.set_font("Helvetica", "B", 15)
        self.cell(0, 8, clean_text(self.school_name), new_x="LMARGIN", new_y="NEXT")

        self.set_x(text_x)
        self.set_font("Helvetica", "", 10)
        label = "Summer Vacation Task" + (" - Answer Key" if self.include_answers else "")
        self.cell(0, 6, label, new_x="LMARGIN", new_y="NEXT")

        self.set_text_color(0, 0, 0)
        self.set_y(28 + max(4, self.side_margin / 2))
        self.set_font("Helvetica", "B", 10)
        self.set_fill_color(*lighten(self.brand_rgb))
        display_topic = self.topic if len(str(self.topic)) <= 70 else str(self.topic)[:67] + "..."
        self.multi_cell(0, 8, clean_text(f"{self.class_level}   |   Subject: {self.subject}   |   Topic: {display_topic}"),
                         fill=True, new_x="LMARGIN", new_y="NEXT")
        self.ln(max(4, self.side_margin / 3))

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(140, 140, 140)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")


def create_pdf(school_name, logo_path, class_level, subject, topic, questions, include_answers, side_margin=15,
                brand_color="2E7D32", margin_mode="uniform", binding_extra=10, show_margin_lines=False):
    import os
    brand_rgb = hex_to_rgb(brand_color)
    light_rgb = lighten(brand_rgb)
    bottom_margin_value = side_margin + 7
    pdf = StyledPDF(school_name, logo_path, class_level, subject, topic, include_answers, brand_color=brand_color,
                     side_margin=side_margin, margin_mode=margin_mode, binding_extra=binding_extra,
                     show_margin_lines=show_margin_lines, bottom_margin=bottom_margin_value)
    pdf.set_auto_page_break(auto=True, margin=bottom_margin_value)

    urdu_available = os.path.exists(URDU_FONT_PATH)
    if urdu_available:
        pdf.add_font("NotoUrdu", "", URDU_FONT_PATH)

    pdf.add_page()
    pdf.set_font("Helvetica", "", 11)

    def print_line(text, height=7, bold=False, italic=False, color=(0, 0, 0)):
        pdf.set_text_color(*color)
        pdf.set_x(pdf.l_margin)
        if is_urdu_text(text) and urdu_available:
            pdf.set_font("NotoUrdu", "", 13)
            pdf.multi_cell(0, height + 2, shape_urdu(text), align="R", new_x="LMARGIN", new_y="NEXT")
        else:
            style = ("B" if bold else "") + ("I" if italic else "")
            pdf.set_font("Helvetica", style, 11)
            pdf.multi_cell(0, height, clean_text(str(text)), new_x="LMARGIN", new_y="NEXT")

    def render_list(q_list):
        for i, q in enumerate(q_list, 1):
            if pdf.get_y() > pdf.page_break_trigger - 25:
                pdf.add_page()

            if q["type"] == "mind_map":
                print_line("Mind Map: " + q["title"], bold=True, color=brand_rgb)
                for line in mindmap_to_pdf_lines(q)[1:]:
                    print_line(line, height=6)
                pdf.ln(2)
                continue

            if q["type"] in ("gender_table", "number_table"):
                print_line(q["title"], bold=True, color=brand_rgb)
                pdf.ln(1)
                col_width = (210 - pdf.l_margin - pdf.r_margin) / 2
                row_h = 10

                if urdu_available:
                    pdf.set_font("NotoUrdu", "", 12)
                    header_word, header_ans = shape_urdu("لفظ"), shape_urdu("جواب")
                else:
                    pdf.set_font("Helvetica", "B", 11)
                    header_word, header_ans = "Word", "Answer"

                pdf.set_fill_color(*light_rgb)
                pdf.set_text_color(0, 0, 0)
                pdf.cell(col_width, row_h, header_ans if include_answers else "", border=1, align="C", fill=True)
                pdf.cell(col_width, row_h, header_word, border=1, align="C", fill=True)
                pdf.ln(row_h)

                for p in q["pairs"]:
                    if urdu_available:
                        pdf.set_font("NotoUrdu", "", 12)
                        word_cell = shape_urdu(p["word"])
                        ans_cell = shape_urdu(p["answer"]) if include_answers else ""
                    else:
                        pdf.set_font("Helvetica", "", 11)
                        word_cell = clean_text(p["word"])
                        ans_cell = clean_text(p["answer"]) if include_answers else ""
                    pdf.cell(col_width, row_h, ans_cell, border=1, align="C")
                    pdf.cell(col_width, row_h, word_cell, border=1, align="C")
                    pdf.ln(row_h)

                pdf.ln(4)
                continue

            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(*brand_rgb)
            pdf.set_x(pdf.l_margin)
            pdf.cell(0, 6, clean_text(f"Q{i}"), new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)

            print_line(q["question"])

            if q["type"] == "mcq":
                for idx, opt in enumerate(q["options"]):
                    print_line(f"     {chr(65 + idx)}.  {opt}", height=6)

            if include_answers:
                print_line(f"Answer: {q['answer']}", height=6, italic=True, color=brand_rgb)
                if "solution" in q:
                    print_line(f"Solution: {q['solution']}", height=6, color=(90, 90, 90))

            pdf.ln(3)
            pdf.set_draw_color(*light_rgb)
            pdf.line(pdf.l_margin, pdf.get_y(), 210 - pdf.r_margin, pdf.get_y())
            pdf.ln(5)

    if isinstance(questions, dict) and "subjects" in questions:
        for block in questions["subjects"]:
            print_line(f"{block['subject']} \u2014 {block['topic']}", height=8, bold=True, color=brand_rgb)
            pdf.ln(1)
            render_list(block["questions"])
            pdf.ln(4)
    else:
        render_list(questions)

    pdf_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
    pdf.output(pdf_path)
    return pdf_path


def create_docx(school_name, logo_path, class_level, subject, topic, questions, include_answers, brand_color="2E7D32"):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    rgb = hex_to_rgb(brand_color)
    color = RGBColor(*rgb)

    doc = Document()

    if logo_path:
        try:
            doc.add_picture(logo_path, width=Inches(0.8))
        except Exception:
            pass

    title = doc.add_heading(school_name, level=1)
    title.runs[0].font.color.rgb = color

    subtitle = doc.add_paragraph("Summer Vacation Task" + (" - Answer Key" if include_answers else ""))
    subtitle.runs[0].font.size = Pt(12)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"{class_level}   |   Subject: {subject}   |   Topic: {topic}")
    meta_run.bold = True
    doc.add_paragraph("")

    def add_question_list(q_list):
        for i, q in enumerate(q_list, 1):
            if q["type"] == "mind_map":
                p = doc.add_paragraph()
                r = p.add_run("Mind Map: " + q["title"])
                r.bold = True
                r.font.color.rgb = color
                for line in mindmap_to_pdf_lines(q)[1:]:
                    doc.add_paragraph(line)
                doc.add_paragraph("")
                continue

            if q["type"] in ("gender_table", "number_table"):
                p = doc.add_paragraph()
                r = p.add_run(q["title"])
                r.bold = True
                r.font.color.rgb = color
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                table = doc.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "جواب" if include_answers else ""
                hdr[1].text = "لفظ"
                for cell in hdr:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for p_item in q["pairs"]:
                    row = table.add_row().cells
                    row[0].text = p_item["answer"] if include_answers else ""
                    row[1].text = p_item["word"]
                    for cell in row:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph("")
                continue

            qp = doc.add_paragraph()
            qnum = qp.add_run(f"Q{i}. ")
            qnum.bold = True
            qnum.font.color.rgb = color
            q_run = qp.add_run(q["question"])
            if is_urdu_text(q["question"]):
                qp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            if q["type"] == "mcq":
                for idx, opt in enumerate(q["options"]):
                    doc.add_paragraph(f"     {chr(65 + idx)}.  {opt}")

            if include_answers:
                ap = doc.add_paragraph()
                ar = ap.add_run(f"Answer: {q['answer']}")
                ar.italic = True
                ar.font.color.rgb = color
                if is_urdu_text(str(q["answer"])):
                    ap.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if "solution" in q:
                    sp = doc.add_paragraph()
                    sr = sp.add_run(f"Solution: {q['solution']}")
                    sr.font.color.rgb = RGBColor(90, 90, 90)

            doc.add_paragraph("")

    if isinstance(questions, dict) and "subjects" in questions:
        for block in questions["subjects"]:
            hp = doc.add_paragraph()
            hr = hp.add_run(f"{block['subject']} \u2014 {block['topic']}")
            hr.bold = True
            hr.font.size = Pt(14)
            hr.font.color.rgb = color
            add_question_list(block["questions"])
    else:
        add_question_list(questions)

    docx_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(docx_path)
    return docx_path


# --- Public shareable view (no login required) ---
query_task_id = st.query_params.get("task")
if query_task_id:
    task_result = db.table("tasks").select("*").eq("id", query_task_id).execute().data
    if not task_result:
        st.error("This task link is invalid or the task was deleted.")
        st.stop()

    t = task_result[0]
    school_result = db.table("school_profile").select("*").eq("id", t["school_id"]).execute().data
    school_name_public = school_result[0]["school_name"] if school_result else "School"

    st.title(school_name_public)
    st.caption(f"Summer Vacation Task — {t['class_level']}, {t['subject']}, {t['topic']}")

    questions = t["questions_json"]
    if isinstance(questions, dict) and "subjects" in questions:
        for block in questions["subjects"]:
            st.markdown(f"**{block['subject']} — {block['topic']}**")
            for i, q in enumerate(block["questions"], 1):
                if q["type"] == "mind_map":
                    st.write(f"**Mind Map: {q['title']}**")
                    markmap(mindmap_to_markdown(q), height=400)
                elif q["type"] in ("gender_table", "number_table"):
                    render_table_card(q, show_answer=False)
                else:
                    render_question_card(i, q, show_answer=False)
    else:
        for i, q in enumerate(questions, 1):
            if q["type"] == "mind_map":
                st.write(f"**Mind Map: {q['title']}**")
                markmap(mindmap_to_markdown(q), height=400)
            elif q["type"] in ("gender_table", "number_table"):
                render_table_card(q, show_answer=False)
            else:
                render_question_card(i, q, show_answer=False)

    if st.button("Download PDF"):
        path = create_pdf(school_name_public, None, t['class_level'], t['subject'], t['topic'], questions, include_answers=False)
        with open(path, "rb") as f:
            st.download_button("Click to save PDF", f, file_name="student_task.pdf", mime="application/pdf")

    st.stop()  # public visitors stop here — never see the teacher login screen


# --- Session state ---
if "user" not in st.session_state:
    st.session_state["user"] = None


# --- AUTH SCREEN ---
if not st.session_state["user"]:
    st.title("Summer Task Generator — Teacher Login")
    tab1, tab2, tab3 = st.tabs(["Log In", "Sign Up", "Forgot Password"])

    with tab1:
        email = st.text_input("Email", key="login_email")
        password = st.text_input("Password", type="password", key="login_password")
        if st.button("Log In"):
            try:
                result = sign_in(email, password)
                st.session_state["user"] = result.user
                st.session_state["access_token"] = result.session.access_token
                st.session_state["refresh_token"] = result.session.refresh_token
                st.rerun()
            except Exception as e:
                st.error(f"Login failed: {e}")

    with tab2:
        new_email = st.text_input("Email", key="signup_email")
        new_password = st.text_input("Password", type="password", key="signup_password")
        if st.button("Sign Up"):
            try:
                sign_up(new_email, new_password)
                st.success("Account created! Check your email to confirm, then log in.")
            except Exception as e:
                st.error(f"Sign up failed: {e}")

    with tab3:
        reset_email = st.text_input("Email", key="reset_email")
        if st.button("Send Reset Link"):
            try:
                reset_password(reset_email)
                st.success("Password reset email sent.")
            except Exception as e:
                st.error(f"Reset failed: {e}")

    st.stop()


# --- LOGGED IN ---
# Re-attach the saved login session to this fresh connection so
# Supabase's Row Level Security recognizes us as authenticated.
user = st.session_state["user"]
st.sidebar.write(f"Logged in as: {user.email}")
if st.sidebar.button("Log Out"):
    sign_out()
    st.rerun()

# Get or create teacher's school profile
teacher_row = db.table("teachers").select("*").eq("id", user.id).execute()

if not teacher_row.data:
    st.title("Welcome! Set Up or Join a School")
    onboard_tab1, onboard_tab2 = st.tabs(["Create New School", "Join Existing School"])

    with onboard_tab1:
        school_name_input = st.text_input("School Name", key="new_school_name")
        if st.button("Create School"):
            code = pysecrets.token_hex(3).upper()
            school_result = db.table("school_profile").insert({"school_name": school_name_input, "school_code": code}).execute()
            new_school_id = school_result.data[0]["id"]
            db.table("teachers").insert({
                "id": user.id, "school_id": new_school_id, "full_name": user.email, "role": "admin"
            }).execute()
            st.rerun()

    with onboard_tab2:
        st.write("Ask your school's admin for the School Code, then enter it below.")
        join_code = st.text_input("School Code", key="join_code")
        if st.button("Join School"):
            match = db.table("school_profile").select("*").eq("school_code", join_code.strip().upper()).execute().data
            if not match:
                st.error("No school found with that code. Please check and try again.")
            else:
                db.table("teachers").insert({
                    "id": user.id, "school_id": match[0]["id"], "full_name": user.email, "role": "teacher"
                }).execute()
                st.rerun()

    st.stop()

school_id = teacher_row.data[0]["school_id"]
my_role = teacher_row.data[0]["role"]
school = db.table("school_profile").select("*").eq("id", school_id).execute().data[0]
school_name = school["school_name"]

if not school.get("school_code"):
    new_code = pysecrets.token_hex(3).upper()
    db.table("school_profile").update({"school_code": new_code}).eq("id", school_id).execute()
    school["school_code"] = new_code

brand_color = school.get("brand_color") or "2E7D32"

st.markdown(f"""
<div class="app-header">
    <h1>📚 Summer Task Generator</h1>
    <p>{school_name}</p>
</div>
""", unsafe_allow_html=True)

with st.sidebar:
    st.markdown(f"### {school_name}")
    st.caption(f"👤 {user.email}")
    st.write("")

    nav_options = ["Dashboard", "Create Task", "Task History", "Question Bank"]
    nav_icons = ["speedometer2", "plus-circle", "clock-history", "collection"]
    if my_role == "admin":
        nav_options.append("School Settings")
        nav_icons.append("gear")

    selected_page = option_menu(
        menu_title=None,
        options=nav_options,
        icons=nav_icons,
        default_index=0,
        styles={
            "container": {"padding": "0", "background-color": "transparent"},
            "icon": {"color": f"#{brand_color}", "font-size": "16px"},
            "nav-link": {"font-size": "15px", "text-align": "left", "margin": "2px 0", "border-radius": "8px"},
            "nav-link-selected": {"background-color": f"#{brand_color}"},
        }
    )

if selected_page == "Dashboard":
    all_school_tasks_dash = db.table("tasks").select("*").eq("school_id", school_id).order("created_at", desc=True).execute().data

    st.markdown(f"""
    <div class="dash-welcome">
        <h2>Welcome back 👋</h2>
        <p>Here's what's happening at {school_name}.</p>
    </div>
    """, unsafe_allow_html=True)

    from datetime import datetime
    this_month = datetime.utcnow().strftime("%Y-%m")
    tasks_this_month = sum(1 for t in all_school_tasks_dash if t["created_at"].startswith(this_month))

    subject_counts = {}
    for t in all_school_tasks_dash:
        subject_counts[t["subject"]] = subject_counts.get(t["subject"], 0) + 1
    top_subject = max(subject_counts, key=subject_counts.get) if subject_counts else "—"

    daily_limit_dash = school.get("daily_generation_limit") or 60
    _, used_today_dash, limit_dash = check_quota(school_id, daily_limit_dash)

    col_s1, col_s2, col_s3, col_s4 = st.columns(4)
    with col_s1:
        st.markdown(f'<div class="stat-card"><div class="stat-number">{len(all_school_tasks_dash)}</div><div class="stat-label">Total Tasks</div></div>', unsafe_allow_html=True)
    with col_s2:
        st.markdown(f'<div class="stat-card"><div class="stat-number">{tasks_this_month}</div><div class="stat-label">Tasks This Month</div></div>', unsafe_allow_html=True)
    with col_s3:
        st.markdown(f'<div class="stat-card"><div class="stat-number">{top_subject}</div><div class="stat-label">Most-Used Subject</div></div>', unsafe_allow_html=True)
    with col_s4:
        st.markdown(f'<div class="stat-card"><div class="stat-number">{used_today_dash}/{limit_dash}</div><div class="stat-label">AI Generations Today</div></div>', unsafe_allow_html=True)

    st.write("")
    st.write("")
    st.subheader("Quick Actions")
    qa1, qa2, qa3 = st.columns(3)
    with qa1:
        st.info("**➕ Create New Task**\n\nGenerate a fresh set of questions.")
    with qa2:
        st.info("**📚 Question Bank**\n\nReuse saved questions in a new task.")
    with qa3:
        st.info("**🕘 Task History**\n\nRevisit and re-download past tasks.")
    st.caption("Use the sidebar to jump to any of these.")

    st.write("")
    st.subheader("Recent Tasks")
    if not all_school_tasks_dash:
        st.write("No tasks created yet — head to **Create Task** in the sidebar to get started.")
    else:
        for t in all_school_tasks_dash[:5]:
            st.markdown(
                f'<div class="recent-task-row">📄 <b>{t["subject"]}</b> — {t["topic"]} '
                f'&nbsp;·&nbsp; {t["class_level"]} &nbsp;·&nbsp; <span style="color:#888;">{t["created_at"][:10]}</span></div>',
                unsafe_allow_html=True
            )

if selected_page == "Create Task":
    logo_file = st.file_uploader("School Logo (optional, per PDF)", type=["png", "jpg", "jpeg"])
    logo_path = None
    if logo_file:
        logo_path = tempfile.NamedTemporaryFile(delete=False, suffix=".png").name
        with open(logo_path, "wb") as f:
            f.write(logo_file.getvalue())

    task_mode = st.radio("Task Mode", ["Single Subject", "Multi-Subject Booklet"])

    # ============== SINGLE SUBJECT MODE ==============
    if task_mode == "Single Subject":
        st.subheader("Task Details")
        class_level = st.selectbox("Class", CLASS_OPTIONS, index=CLASS_OPTIONS.index("Class 7"))
        subject = st.text_input("Subject", "Science")

        source_choice = st.radio("Question source", ["Type a topic", "Upload notes/chapter file"])

        topic = ""
        source_text = None
        chapter_focus = None

        if source_choice == "Type a topic":
            topic = st.text_input("Topic", "Photosynthesis")
        else:
            notes_file = st.file_uploader("Upload notes/chapter/book (PDF, DOCX, or TXT)", type=["pdf", "docx", "txt"])
            if notes_file:
                page_start, page_end = None, None
                if notes_file.name.lower().endswith(".pdf"):
                    total_pages = get_pdf_page_count(notes_file)
                    st.caption(f"This PDF has {total_pages} pages.")
                    use_range = st.checkbox("Only use a specific page range (e.g. one chapter of a book)")
                    if use_range:
                        col_p1, col_p2 = st.columns(2)
                        with col_p1:
                            page_start = st.number_input("From page", min_value=1, max_value=total_pages, value=1)
                        with col_p2:
                            page_end = st.number_input("To page", min_value=1, max_value=total_pages, value=min(5, total_pages))

                chapter_focus = st.text_input("Focus on a specific chapter/topic within this file (optional)", "")

                with st.spinner("Reading file..."):
                    source_text = extract_text_from_file(notes_file, page_start, page_end)
                if source_text:
                    st.success(f"Extracted {len(source_text)} characters" + (f" from pages {page_start}-{page_end}" if page_start else "") + ".")
                    if chapter_focus:
                        topic = chapter_focus
                    else:
                        clean_name = notes_file.name.rsplit(".", 1)[0]
                        topic = clean_name if len(clean_name) <= 50 else clean_name[:47] + "..."
                else:
                    st.error("Couldn't extract text from this file (or this page range). Try different pages or a different file.")

        st.write("**Question Types** (select one or more, mix and match)")
        selected_types = st.multiselect(
            "Question Types",
            ALL_QUESTION_TYPES,
            default=["mcq"],
            format_func=lambda t: TYPE_LABELS.get(t, t),
            label_visibility="collapsed"
        )
        if any(t in URDU_SCRIPT_TYPES for t in selected_types):
            st.caption("✅ Urdu Grammar / Translation questions now render correctly in the PDF too (right-to-left Urdu script support).")

        type_counts = {}
        if selected_types:
            cols = st.columns(len(selected_types))
            for idx, qt in enumerate(selected_types):
                with cols[idx]:
                    type_counts[qt] = st.number_input(TYPE_LABELS.get(qt, qt), min_value=1, max_value=20, value=5, key=f"count_{qt}")

        col_d1, col_d2 = st.columns(2)
        with col_d1:
            difficulty = st.selectbox("Difficulty", ["easy", "medium", "hard"])
        with col_d2:
            answer_length = st.selectbox("Answer Length (for theory/essay-type questions)", ["short", "medium", "long"], index=1)

        if st.button("Generate Questions"):
            daily_limit = school.get("daily_generation_limit") or 60
            allowed, used_today, limit = check_quota(school_id, daily_limit)
            if not selected_types:
                st.error("Please select at least one question type.")
            elif source_choice == "Upload notes/chapter file" and not source_text:
                st.error("Please upload a valid file first.")
            elif not allowed:
                st.error(f"Your school has reached its daily AI generation limit ({used_today}/{limit}). This protects the shared free AI quota \u2014 it resets tomorrow, or your admin can raise the limit in School Settings.")
            else:
                all_questions = []
                with st.spinner("Generating..."):
                    for qt in selected_types:
                        qs = generate_questions(class_level, subject, topic, qt, type_counts[qt], difficulty,
                                                 source_text=source_text, chapter_focus=chapter_focus, answer_length=answer_length)
                        all_questions.extend(qs)
                log_usage(school_id, user.id)
                st.session_state["questions"] = all_questions
                st.session_state["task_meta"] = {
                    "class_level": class_level, "subject": subject, "topic": topic, "difficulty": difficulty
                }

                overall_type = selected_types[0] if len(selected_types) == 1 else "mixed"
                db.table("tasks").insert({
                    "school_id": school_id,
                    "created_by": user.id,
                    "class_level": class_level,
                    "subject": subject,
                    "topic": topic,
                    "question_type": overall_type,
                    "difficulty": difficulty,
                    "questions_json": all_questions
                }).execute()

        if "questions" in st.session_state:
            meta = st.session_state.get("task_meta", {})
            st.subheader("Preview")
            for i, q in enumerate(st.session_state["questions"], 1):
                if q["type"] == "mind_map":
                    st.write(f"**Mind Map: {q['title']}**")
                    markmap(mindmap_to_markdown(q), height=400)
                elif q["type"] in ("gender_table", "number_table"):
                    render_table_card(q, show_answer=True)
                else:
                    render_question_card(i, q, show_answer=True)

                    col_e1, col_e2 = st.columns([1, 1])
                    with col_e1:
                        if st.button("🔄 Regenerate", key=f"regen_{i}"):
                            daily_limit = school.get("daily_generation_limit") or 60
                            allowed, used_today, limit = check_quota(school_id, daily_limit)
                            if not allowed:
                                st.error(f"Daily AI generation limit reached ({used_today}/{limit}) for your school.")
                            else:
                                with st.spinner("Regenerating..."):
                                    new_q = generate_questions(
                                        meta.get("class_level", class_level), meta.get("subject", subject),
                                        meta.get("topic", topic), q["type"], 1, meta.get("difficulty", difficulty)
                                    )
                                log_usage(school_id, user.id)
                                if new_q:
                                    st.session_state["questions"][i - 1] = new_q[0]
                                else:
                                    st.warning("The AI couldn't produce a valid replacement question \u2014 try again.")
                                st.rerun()
                    with col_e2:
                        with st.expander(f"✏️ Edit Question {i}"):
                            edited_text = st.text_area("Question text", q["question"], key=f"edit_q_{i}")
                            edited_answer = st.text_input("Answer", q["answer"], key=f"edit_a_{i}")
                            edited_options_raw = None
                            if q["type"] == "mcq":
                                edited_options_raw = st.text_area("Options (one per line)", "\n".join(q["options"]), key=f"edit_o_{i}")
                            if st.button("Save Edit", key=f"save_edit_{i}"):
                                st.session_state["questions"][i - 1]["question"] = edited_text
                                st.session_state["questions"][i - 1]["answer"] = edited_answer
                                if edited_options_raw is not None:
                                    st.session_state["questions"][i - 1]["options"] = [o.strip() for o in edited_options_raw.split("\n") if o.strip()]
                                st.success("Saved.")
                                st.rerun()

            if st.button("💾 Save these questions to Question Bank"):
                rows = [{
                    "school_id": school_id,
                    "class_level": meta.get("class_level", class_level),
                    "subject": meta.get("subject", subject),
                    "topic": meta.get("topic", topic),
                    "question_type": q["type"],
                    "question_data": q,
                    "created_by": user.id
                } for q in st.session_state["questions"] if q["type"] not in ("mind_map", "gender_table", "number_table")]
                if rows:
                    db.table("question_bank").insert(rows).execute()
                    st.success(f"Saved {len(rows)} question(s) to the Question Bank.")
                else:
                    st.warning("Mind maps can't be saved to the Question Bank individually.")

            with st.expander("✏️ Edit Task Info (class / subject / topic shown on the PDF)"):
                meta_class_default = meta.get("class_level", class_level)
                meta_class_index = CLASS_OPTIONS.index(meta_class_default) if meta_class_default in CLASS_OPTIONS else 0
                edited_class = st.selectbox("Class", CLASS_OPTIONS, index=meta_class_index, key="meta_edit_class")
                edited_subject = st.text_input("Subject", meta.get("subject", subject), key="meta_edit_subject")
                edited_topic = st.text_input("Topic", meta.get("topic", topic), key="meta_edit_topic")
                if st.button("Update Task Info"):
                    st.session_state["task_meta"] = {
                        "class_level": edited_class, "subject": edited_subject,
                        "topic": edited_topic, "difficulty": meta.get("difficulty", difficulty)
                    }
                    st.success("Updated — this now applies to the PDF and Question Bank saving below.")
                    st.rerun()

            col_m1, col_m2 = st.columns([1, 1])
            with col_m1:
                pdf_margin_choice = st.selectbox("PDF/Word Margins", ["Narrow", "Normal", "Wide", "Custom"], index=1)
            if pdf_margin_choice == "Custom":
                with col_m2:
                    pdf_margin_value = st.number_input("Custom margin (mm)", min_value=5, max_value=40, value=15)
            else:
                pdf_margin_value = {"Narrow": 10, "Normal": 15, "Wide": 22}[pdf_margin_choice]

            col_ml1, col_ml2 = st.columns([1, 1])
            with col_ml1:
                margin_layout_choice = st.selectbox(
                    "PDF Margin Layout",
                    ["4-sided (uniform)", "1-sided (binding margin, always left)", "2-sided / duplex (mirrored binding margin)"],
                    key="pdf_margin_layout"
                )
            margin_mode = {"4-sided (uniform)": "uniform", "1-sided (binding margin, always left)": "single",
                            "2-sided / duplex (mirrored binding margin)": "double"}[margin_layout_choice]
            if margin_mode != "uniform":
                with col_ml2:
                    binding_extra_value = st.number_input("Extra binding margin (mm)", min_value=5, max_value=30, value=10, key="pdf_binding_extra")
            else:
                binding_extra_value = 10

            show_margin_lines = st.checkbox("Show margin guide lines on PDF (light gray border showing the margin, useful for printing/cutting reference)", key="pdf_show_margins")

            st.write("**Download as:**")
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                if st.button("Student PDF"):
                    path = create_pdf(school_name, logo_path, meta.get("class_level", class_level),
                                       meta.get("subject", subject), meta.get("topic", topic),
                                       st.session_state["questions"], include_answers=False, side_margin=pdf_margin_value,
                                       brand_color=brand_color, margin_mode=margin_mode, binding_extra=binding_extra_value,
                                       show_margin_lines=show_margin_lines)
                    with open(path, "rb") as f:
                        st.download_button("Download Student PDF", f, file_name="student_task.pdf", mime="application/pdf")
            with col2:
                if st.button("Teacher PDF"):
                    path = create_pdf(school_name, logo_path, meta.get("class_level", class_level),
                                       meta.get("subject", subject), meta.get("topic", topic),
                                       st.session_state["questions"], include_answers=True, side_margin=pdf_margin_value,
                                       brand_color=brand_color, margin_mode=margin_mode, binding_extra=binding_extra_value,
                                       show_margin_lines=show_margin_lines)
                    with open(path, "rb") as f:
                        st.download_button("Download Teacher PDF", f, file_name="teacher_answer_key.pdf", mime="application/pdf")
            with col3:
                if st.button("Student Word"):
                    path = create_docx(school_name, logo_path, meta.get("class_level", class_level),
                                        meta.get("subject", subject), meta.get("topic", topic),
                                        st.session_state["questions"], include_answers=False, brand_color=brand_color)
                    with open(path, "rb") as f:
                        st.download_button("Download Student Word", f, file_name="student_task.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with col4:
                if st.button("Teacher Word"):
                    path = create_docx(school_name, logo_path, meta.get("class_level", class_level),
                                        meta.get("subject", subject), meta.get("topic", topic),
                                        st.session_state["questions"], include_answers=True, brand_color=brand_color)
                    with open(path, "rb") as f:
                        st.download_button("Download Teacher Word", f, file_name="teacher_answer_key.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # ============== MULTI-SUBJECT BOOKLET MODE ==============
    else:
        st.subheader("Booklet Details")
        st.caption("Tip: to get specific question counts per topic within the SAME subject, just repeat the same subject name in each block below (e.g. Subject 1 = 'Math' / Topic 1 = 'LCM', Subject 2 = 'Math' / Topic 2 = 'DMAS') \u2014 each block gets its own question types and counts.")
        booklet_class = st.selectbox("Class", CLASS_OPTIONS, index=CLASS_OPTIONS.index("Class 7"), key="booklet_class")
        num_subjects = st.number_input("How many subject/topic blocks?", min_value=2, max_value=8, value=2)

        subject_blocks = []
        for i in range(int(num_subjects)):
            with st.expander(f"Block {i + 1}", expanded=(i < 2)):
                c1, c2 = st.columns(2)
                with c1:
                    b_subj = st.text_input(f"Subject", key=f"booklet_subject_{i}")
                with c2:
                    b_topic = st.text_input(f"Topic", key=f"booklet_topic_{i}")

                b_types = st.multiselect(
                    "Question Types for this block",
                    ALL_QUESTION_TYPES,
                    default=["mcq"],
                    key=f"booklet_types_{i}",
                    format_func=lambda t: TYPE_LABELS.get(t, t)
                )
                if any(t in URDU_SCRIPT_TYPES for t in b_types):
                    st.caption("✅ Urdu Grammar / Translation questions render correctly in the PDF too.")

                b_counts = {}
                if b_types:
                    cols = st.columns(len(b_types))
                    for idx, qt in enumerate(b_types):
                        with cols[idx]:
                            b_counts[qt] = st.number_input(TYPE_LABELS.get(qt, qt), min_value=1, max_value=20, value=3, key=f"booklet_count_{i}_{qt}")

                subject_blocks.append({"subject": b_subj, "topic": b_topic, "types": b_types, "counts": b_counts})

        col_bd1, col_bd2 = st.columns(2)
        with col_bd1:
            booklet_difficulty = st.selectbox("Difficulty (applies to all blocks)", ["easy", "medium", "hard"], key="booklet_difficulty")
        with col_bd2:
            booklet_answer_length = st.selectbox("Answer Length", ["short", "medium", "long"], index=1, key="booklet_answer_length")

        if st.button("Generate Booklet"):
            daily_limit = school.get("daily_generation_limit") or 60
            allowed, used_today, limit = check_quota(school_id, daily_limit)
            if any(not b["subject"] or not b["topic"] for b in subject_blocks):
                st.error("Please fill in every subject and topic.")
            elif any(not b["types"] for b in subject_blocks):
                st.error("Please select at least one question type for every block.")
            elif not allowed:
                st.error(f"Your school has reached its daily AI generation limit ({used_today}/{limit}). This protects the shared free AI quota \u2014 it resets tomorrow, or your admin can raise the limit in School Settings.")
            else:
                booklet_subjects = []
                with st.spinner("Generating booklet..."):
                    for block in subject_blocks:
                        subj_questions = []
                        for qt in block["types"]:
                            qs = generate_questions(booklet_class, block["subject"], block["topic"], qt, block["counts"][qt], booklet_difficulty, answer_length=booklet_answer_length)
                            subj_questions.extend(qs)
                        booklet_subjects.append({"subject": block["subject"], "topic": block["topic"], "questions": subj_questions})
                log_usage(school_id, user.id)

                booklet_data = {"subjects": booklet_subjects}
                st.session_state["booklet"] = booklet_data
                st.session_state["booklet_class"] = booklet_class

                db.table("tasks").insert({
                    "school_id": school_id,
                    "created_by": user.id,
                    "class_level": booklet_class,
                    "subject": "Multiple Subjects",
                    "topic": "Combined Booklet",
                    "question_type": "booklet",
                    "difficulty": booklet_difficulty,
                    "questions_json": booklet_data
                }).execute()

        if "booklet" in st.session_state:
            st.subheader("Booklet Preview")
            for block in st.session_state["booklet"]["subjects"]:
                st.markdown(f"### {block['subject']} — {block['topic']}")
                for i, q in enumerate(block["questions"], 1):
                    if q["type"] == "mind_map":
                        st.write(f"**Mind Map: {q['title']}**")
                        markmap(mindmap_to_markdown(q), height=400)
                    elif q["type"] in ("gender_table", "number_table"):
                        render_table_card(q, show_answer=True)
                    else:
                        render_question_card(i, q, show_answer=True)

            col_bm1, col_bm2 = st.columns([1, 1])
            with col_bm1:
                booklet_margin_choice = st.selectbox("PDF/Word Margins", ["Narrow", "Normal", "Wide", "Custom"], index=1, key="booklet_margin")
            if booklet_margin_choice == "Custom":
                with col_bm2:
                    booklet_margin_value = st.number_input("Custom margin (mm)", min_value=5, max_value=40, value=15, key="booklet_margin_custom")
            else:
                booklet_margin_value = {"Narrow": 10, "Normal": 15, "Wide": 22}[booklet_margin_choice]

            col_bml1, col_bml2 = st.columns([1, 1])
            with col_bml1:
                booklet_margin_layout_choice = st.selectbox(
                    "PDF Margin Layout",
                    ["4-sided (uniform)", "1-sided (binding margin, always left)", "2-sided / duplex (mirrored binding margin)"],
                    key="booklet_margin_layout"
                )
            booklet_margin_mode = {"4-sided (uniform)": "uniform", "1-sided (binding margin, always left)": "single",
                                    "2-sided / duplex (mirrored binding margin)": "double"}[booklet_margin_layout_choice]
            if booklet_margin_mode != "uniform":
                with col_bml2:
                    booklet_binding_extra = st.number_input("Extra binding margin (mm)", min_value=5, max_value=30, value=10, key="booklet_binding_extra")
            else:
                booklet_binding_extra = 10

            booklet_show_margin_lines = st.checkbox("Show margin guide lines on PDF (light gray border showing the margin, useful for printing/cutting reference)", key="booklet_show_margins")

            st.write("**Download as:**")
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                if st.button("Student PDF", key="booklet_student_pdf"):
                    path = create_pdf(school_name, logo_path, st.session_state["booklet_class"], "Multiple Subjects", "Combined Booklet", st.session_state["booklet"], include_answers=False, side_margin=booklet_margin_value, brand_color=brand_color, margin_mode=booklet_margin_mode, binding_extra=booklet_binding_extra, show_margin_lines=booklet_show_margin_lines)
                    with open(path, "rb") as f:
                        st.download_button("Download Student Booklet PDF", f, file_name="student_booklet.pdf", mime="application/pdf")
            with col2:
                if st.button("Teacher PDF", key="booklet_teacher_pdf"):
                    path = create_pdf(school_name, logo_path, st.session_state["booklet_class"], "Multiple Subjects", "Combined Booklet", st.session_state["booklet"], include_answers=True, side_margin=booklet_margin_value, brand_color=brand_color, margin_mode=booklet_margin_mode, binding_extra=booklet_binding_extra, show_margin_lines=booklet_show_margin_lines)
                    with open(path, "rb") as f:
                        st.download_button("Download Teacher Booklet PDF", f, file_name="teacher_booklet.pdf", mime="application/pdf")
            with col3:
                if st.button("Student Word", key="booklet_student_docx"):
                    path = create_docx(school_name, logo_path, st.session_state["booklet_class"], "Multiple Subjects", "Combined Booklet", st.session_state["booklet"], include_answers=False, brand_color=brand_color)
                    with open(path, "rb") as f:
                        st.download_button("Download Student Booklet Word", f, file_name="student_booklet.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with col4:
                if st.button("Teacher Word", key="booklet_teacher_docx"):
                    path = create_docx(school_name, logo_path, st.session_state["booklet_class"], "Multiple Subjects", "Combined Booklet", st.session_state["booklet"], include_answers=True, brand_color=brand_color)
                    with open(path, "rb") as f:
                        st.download_button("Download Teacher Booklet Word", f, file_name="teacher_booklet.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if selected_page == "Task History":
    st.subheader("Past Tasks")
    past_tasks = db.table("tasks").select("*").eq("school_id", school_id).order("created_at", desc=True).execute().data

    if not past_tasks:
        st.write("No tasks created yet.")
    else:
        all_teachers_lookup = {tch["id"]: tch["full_name"] for tch in db.table("teachers").select("*").eq("school_id", school_id).execute().data}

        import csv
        import io
        csv_buffer = io.StringIO()
        writer = csv.writer(csv_buffer)
        writer.writerow(["Date", "Class", "Subject", "Topic", "Question Type", "Difficulty", "Created By", "Question Count"])
        for t in past_tasks:
            qdata = t["questions_json"]
            if isinstance(qdata, dict) and "subjects" in qdata:
                qcount = sum(len(b["questions"]) for b in qdata["subjects"])
            else:
                qcount = len(qdata) if isinstance(qdata, list) else 0
            writer.writerow([
                "'" + t["created_at"][:10], t["class_level"], t["subject"], t["topic"],
                t["question_type"], t["difficulty"], all_teachers_lookup.get(t.get("created_by"), "Unknown"), qcount
            ])
        st.download_button("📊 Export Task History as CSV", csv_buffer.getvalue(), file_name="task_history.csv", mime="text/csv")

        for t in past_tasks:
            with st.expander(f"{t['subject']} - {t['topic']} ({t['class_level']}, {t['question_type']}) — {t['created_at'][:10]}"):
                share_url = f"{app_url}/?task={t['id']}" if app_url else f"?task={t['id']}"
                st.text_input("Shareable link for students/parents", value=share_url, key=f"share_{t['id']}")
                if not app_url:
                    st.caption("Tip: add APP_URL in Streamlit secrets (your app's live URL) to get a full clickable link.")

                questions = t["questions_json"]
                if isinstance(questions, dict) and "subjects" in questions:
                    for block in questions["subjects"]:
                        st.markdown(f"**{block['subject']} — {block['topic']}**")
                        for i, q in enumerate(block["questions"], 1):
                            if q["type"] == "mind_map":
                                st.write(f"**Mind Map: {q['title']}**")
                                markmap(mindmap_to_markdown(q), height=400)
                            elif q["type"] in ("gender_table", "number_table"):
                                render_table_card(q, show_answer=True)
                            else:
                                render_question_card(i, q, show_answer=True)
                else:
                    for i, q in enumerate(questions, 1):
                        if q["type"] == "mind_map":
                            st.write(f"**Mind Map: {q['title']}**")
                            markmap(mindmap_to_markdown(q), height=400)
                        elif q["type"] in ("gender_table", "number_table"):
                            render_table_card(q, show_answer=True)
                        else:
                            render_question_card(i, q, show_answer=True)

                colh1, colh2 = st.columns(2)
                with colh1:
                    if st.button("Download Student PDF", key=f"student_{t['id']}"):
                        path = create_pdf(school_name, None, t['class_level'], t['subject'], t['topic'], questions, include_answers=False, brand_color=brand_color)
                        with open(path, "rb") as f:
                            st.download_button("Click to save", f, file_name="student_task.pdf", mime="application/pdf", key=f"dl_student_{t['id']}")
                with colh2:
                    if st.button("Download Teacher PDF", key=f"teacher_{t['id']}"):
                        path = create_pdf(school_name, None, t['class_level'], t['subject'], t['topic'], questions, include_answers=True, brand_color=brand_color)
                        with open(path, "rb") as f:
                            st.download_button("Click to save", f, file_name="teacher_answer_key.pdf", mime="application/pdf", key=f"dl_teacher_{t['id']}")

                with st.expander("✏️ Rename this task"):
                    rn_class = st.selectbox("Class", CLASS_OPTIONS, index=CLASS_OPTIONS.index(t['class_level']) if t['class_level'] in CLASS_OPTIONS else 0, key=f"rn_class_{t['id']}")
                    rn_subject = st.text_input("Subject", t['subject'], key=f"rn_subject_{t['id']}")
                    rn_topic = st.text_input("Topic", t['topic'], key=f"rn_topic_{t['id']}")
                    if st.button("Save Changes", key=f"rn_save_{t['id']}"):
                        db.table("tasks").update({
                            "class_level": rn_class, "subject": rn_subject, "topic": rn_topic
                        }).eq("id", t["id"]).execute()
                        st.success("Updated.")
                        st.rerun()

                st.write("")
                confirm_key = f"confirm_delete_{t['id']}"
                if not st.session_state.get(confirm_key):
                    if st.button("🗑️ Delete this task", key=f"del_{t['id']}"):
                        st.session_state[confirm_key] = True
                        st.rerun()
                else:
                    st.warning("Delete this task permanently? This cannot be undone.")
                    col_confirm1, col_confirm2 = st.columns(2)
                    with col_confirm1:
                        if st.button("Yes, delete it", key=f"del_yes_{t['id']}"):
                            db.table("tasks").delete().eq("id", t["id"]).execute()
                            st.session_state.pop(confirm_key, None)
                            st.success("Task deleted.")
                            st.rerun()
                    with col_confirm2:
                        if st.button("Cancel", key=f"del_no_{t['id']}"):
                            st.session_state.pop(confirm_key, None)
                            st.rerun()

if selected_page == "Question Bank":
    st.subheader("Question Bank")
    st.caption("Reuse questions saved from past tasks to build a new one, without generating fresh questions.")

    bank_items = db.table("question_bank").select("*").eq("school_id", school_id).order("created_at", desc=True).execute().data

    if not bank_items:
        st.write("No saved questions yet. Generate a task and click 'Save to Question Bank' to start building your library.")
    else:
        col_f1, col_f2 = st.columns(2)
        with col_f1:
            filter_class = st.text_input("Filter by Class (optional)", key="bank_filter_class")
        with col_f2:
            filter_subject = st.text_input("Filter by Subject (optional)", key="bank_filter_subject")

        filtered = bank_items
        if filter_class:
            filtered = [b for b in filtered if filter_class.strip().lower() in (b["class_level"] or "").lower()]
        if filter_subject:
            filtered = [b for b in filtered if filter_subject.strip().lower() in (b["subject"] or "").lower()]

        st.write(f"Showing {len(filtered)} of {len(bank_items)} saved questions")

        selected_ids = []
        for b in filtered:
            q = b["question_data"]
            label = f"[{b['class_level']} | {b['subject']} | {b['topic']}] {q['question']}"
            checked = st.checkbox(label, key=f"bank_{b['id']}")
            if checked:
                selected_ids.append(b["id"])

        st.write(f"**{len(selected_ids)} question(s) selected from bank**")
        bundle_class = st.selectbox("Class for this task", CLASS_OPTIONS, index=CLASS_OPTIONS.index("Class 7"), key="bundle_class")
        bundle_subject = st.text_input("Subject for this task", "Mixed", key="bundle_subject")
        bundle_topic = st.text_input("Topic/Title for this task", "Custom Task", key="bundle_topic")

        with st.expander("➕ Also generate new questions to add to this task (optional)"):
            add_new_types = st.multiselect(
                "Additional Question Types",
                ALL_QUESTION_TYPES,
                default=[],
                format_func=lambda t: TYPE_LABELS.get(t, t),
                key="bank_add_types"
            )
            add_new_counts = {}
            add_new_difficulty = "medium"
            if add_new_types:
                cols = st.columns(len(add_new_types))
                for idx, qt in enumerate(add_new_types):
                    with cols[idx]:
                        add_new_counts[qt] = st.number_input(TYPE_LABELS.get(qt, qt), min_value=1, max_value=20, value=3, key=f"bank_add_count_{qt}")
                add_new_difficulty = st.selectbox("Difficulty for new questions", ["easy", "medium", "hard"], index=1, key="bank_new_difficulty")

        if st.button("Build Task"):
            daily_limit = school.get("daily_generation_limit") or 60
            allowed, used_today, limit = check_quota(school_id, daily_limit)
            if not selected_ids and not add_new_types:
                st.error("Select at least one saved question, or add a new question type to generate.")
            elif add_new_types and not allowed:
                st.error(f"Your school has reached its daily AI generation limit ({used_today}/{limit}). You can still build a task from selected saved questions only, or wait until tomorrow.")
            else:
                chosen_questions = [b["question_data"] for b in bank_items if b["id"] in selected_ids]

                if add_new_types:
                    with st.spinner("Generating additional questions..."):
                        for qt in add_new_types:
                            qs = generate_questions(bundle_class, bundle_subject, bundle_topic, qt, add_new_counts[qt], add_new_difficulty)
                            chosen_questions.extend(qs)
                    log_usage(school_id, user.id)

                st.session_state["bank_bundle"] = {
                    "class_level": bundle_class, "subject": bundle_subject,
                    "topic": bundle_topic, "questions": chosen_questions
                }
                db.table("tasks").insert({
                    "school_id": school_id, "created_by": user.id,
                    "class_level": bundle_class, "subject": bundle_subject, "topic": bundle_topic,
                    "question_type": "mixed", "difficulty": "mixed",
                    "questions_json": chosen_questions
                }).execute()
                st.success(f"Task created with {len(chosen_questions)} question(s)! Find it in Task History to download PDFs.")

if selected_page == "School Settings" and my_role == "admin":
        st.subheader("School Code")
        st.write("Share this code with other teachers so they can join your school:")
        st.code(school["school_code"])

        st.subheader("Brand Color")
        st.write("This color is used on your PDF/Word task headers and accents.")
        new_brand_color = st.color_picker("Pick a color", f"#{brand_color}")
        if st.button("Save Brand Color"):
            hex_value = new_brand_color.lstrip("#").upper()
            db.table("school_profile").update({"brand_color": hex_value}).eq("id", school_id).execute()
            st.success("Brand color updated! It will apply to new PDFs/Word docs from now on.")

        st.subheader("Daily AI Generation Limit")
        st.write("The AI service is a shared free resource across all schools using this app. This cap protects against one school accidentally using it all up in a day \u2014 raise it if your school genuinely needs more.")
        current_limit = school.get("daily_generation_limit") or 60
        _, used_today_settings, _ = check_quota(school_id, current_limit)
        st.caption(f"Used today: {used_today_settings} / {current_limit}")
        new_limit = st.number_input("Daily limit (generations per day, school-wide)", min_value=5, max_value=500, value=current_limit)
        if st.button("Save Limit"):
            db.table("school_profile").update({"daily_generation_limit": int(new_limit)}).eq("id", school_id).execute()
            st.success("Daily limit updated.")

        st.subheader("Teachers in this school")
        all_teachers = db.table("teachers").select("*").eq("school_id", school_id).execute().data
        admin_count = sum(1 for tch in all_teachers if tch["role"] == "admin")

        for tch in all_teachers:
            role_label = "Admin" if tch["role"] == "admin" else "Teacher"
            col_t1, col_t2 = st.columns([4, 1])
            with col_t1:
                st.write(f"- {tch['full_name']}  ({role_label})")
            with col_t2:
                is_self = tch["id"] == user.id
                is_last_admin = tch["role"] == "admin" and admin_count <= 1
                if not is_self and not is_last_admin:
                    remove_key = f"confirm_remove_{tch['id']}"
                    if not st.session_state.get(remove_key):
                        if st.button("Remove", key=f"remove_btn_{tch['id']}"):
                            st.session_state[remove_key] = True
                            st.rerun()
                    else:
                        if st.button("Confirm?", key=f"remove_confirm_{tch['id']}"):
                            db.table("teachers").delete().eq("id", tch["id"]).execute()
                            st.session_state.pop(remove_key, None)
                            st.success(f"Removed {tch['full_name']}.")
                            st.rerun()

        st.subheader("Usage Overview")
        all_school_tasks = db.table("tasks").select("*").eq("school_id", school_id).execute().data
        st.write(f"Total tasks created: **{len(all_school_tasks)}**")

        if all_school_tasks:
            counts_by_teacher = {}
            counts_by_subject = {}
            teacher_lookup = {tch["id"]: tch["full_name"] for tch in all_teachers}

            for task in all_school_tasks:
                creator = teacher_lookup.get(task.get("created_by"), "Unknown")
                counts_by_teacher[creator] = counts_by_teacher.get(creator, 0) + 1
                counts_by_subject[task["subject"]] = counts_by_subject.get(task["subject"], 0) + 1

            st.write("**Tasks by teacher:**")
            for name, count in counts_by_teacher.items():
                st.write(f"- {name}: {count}")

            st.write("**Tasks by subject:**")
            for subj, count in counts_by_subject.items():
                st.write(f"- {subj}: {count}")
